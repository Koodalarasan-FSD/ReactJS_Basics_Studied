from flask import Flask, render_template, request
from pymongo import MongoClient
from docx import Document
import os

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/get_properties', methods=['POST'])
def get_properties():
    try:
        MONGODB_LINK = request.form['mongodblink']
        MONGODB_DATABASENAME = request.form['mongodb_databasename']
        MONGODB_COLLECTION_NAME = request.form['mongodb_Collection_name']

        # Connect to MongoDB
        client = MongoClient(MONGODB_LINK)
        db = client[MONGODB_DATABASENAME]
        collection = db[MONGODB_COLLECTION_NAME]

        # Retrieve data from MongoDB
        data = collection.find()

        # Create a new Word document
        doc = Document()

        # Iterate through MongoDB data and write to the document
        for record in data:
            # Customize based on your data structure
            doc.add_paragraph(f"ID: {record['_id']}")
            doc.add_paragraph(f"Name: {record['name']}")
            doc.add_paragraph(f"Type: {record['type']}")
            doc.add_paragraph("-----------------------------")

        # Specify the full path for saving the document
        output_path = os.path.join(os.getcwd(), 'output_document.docx')

        # Save the document
        doc.save(output_path)

        # Inform the user about the successful operation
        return f"Success: Document created successfully at {output_path}"

    except Exception as e:
        # Handle exceptions and display an error message
        error_message = f"Error: {str(e)}"
        return render_template('error.html', error_message=error_message)

if __name__ == "__main__":
    app.run(debug=True)
